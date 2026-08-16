#!/usr/bin/env python3
# ruff: noqa
"""KOS Indexer Format Handlers — 文件格式处理器基类 + Markdown/Docx.

从 indexer/engine.py 抽出 (God Module 拆 wave 7, engine.py 1184->~900).
FormatHandler 基类 + Markdown/Docx handler. Pdf/Xlsx/GongwenDocx 留 engine.py (后续 wave).
"""

from __future__ import annotations

import re
from pathlib import Path


class FormatHandler:
    """文件格式处理器基类 (子类按扩展名路由)."""

    extensions: set[str] = set()
    priority: int = 0

    def can_handle(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.extensions

    def extract_text(self, file_path: Path) -> str:
        raise NotImplementedError

    def extract_metadata(self, file_path: Path) -> dict:  # type: ignore[type-arg]
        return {"title": file_path.stem}


class MarkdownHandler(FormatHandler):
    extensions = {".md", ".txt", ".markdown", ".mdx"}
    priority = 100

    def extract_text(self, p: Path) -> str:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            return p.read_text(encoding="utf-8")[:8000]
        except Exception:  # noqa: BLE001
            return f"[{p.suffix}: {p.name}]"

    def extract_metadata(self, p: Path) -> dict:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            content = p.read_text(encoding="utf-8")
        except Exception:  # noqa: BLE001
            return {"title": p.stem}
        title = p.stem
        m = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
        if m:
            title = m.group(1).strip()
        # Extract KOS frontmatter comment
        fm: dict[str, str] = {}
        km = re.search(r"<!--\s*KOS:\s*(.+?)\s*-->", content[:500])
        if km:
            for part in km.group(1).split():
                if "=" in part:
                    k, v = part.split("=", 1)
                    fm[k.strip()] = v.strip()
        return {"title": title, **fm}


class DocxHandler(FormatHandler):
    extensions = {".docx"}
    priority = 80

    def extract_text(self, p: Path) -> str:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            from docx import Document

            doc = Document(str(p))
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())[:8000]
        except Exception:  # noqa: BLE001
            return f"[docx: {p.name}]"

    def extract_metadata(self, p: Path) -> dict:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            from docx import Document

            doc = Document(str(p))
            return {"title": doc.core_properties.title or p.stem}
        except Exception:  # noqa: BLE001
            return {"title": p.stem}


# 注册表 (供 engine 查询可用 handler)
SIMPLE_HANDLERS: list[FormatHandler] = [MarkdownHandler(), DocxHandler()]
