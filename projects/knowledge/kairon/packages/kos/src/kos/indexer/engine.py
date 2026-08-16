#!/usr/bin/env python3
# ruff: noqa
from kos.config import get_vault_ops_dir  # type: ignore[import-not-found]

"""
KOS Indexer — Fingerprint-driven incremental indexing (v1.0)

Replaces manual full-rebuild with SHA-256 fingerprint diff. Supports
.md, .txt, .docx, .pdf, .xlsx via pluggable FormatHandlers.

Usage:
    python3 kos-indexer.py index                    # Full rebuild
    python3 kos-indexer.py index --incremental      # Incremental (changed only)
    python3 kos-indexer.py index --domain gongwen   # Single domain
    python3 kos-indexer.py index --jobs 8           # Parallel workers (default: 4)
    python3 kos-indexer.py status                   # Index stats
    python3 kos-indexer.py diff                     # Files changed since last index
"""

import hashlib
import fnmatch
import json
import os
import re
import sqlite3
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, cast

from kos.config import get_vault_ops_dir  # type: ignore[[, unused-ignore]

VAULT_OPS_DIR = get_vault_ops_dir()  # type: ignore[func-returns-value]

from kos._default_workspace_config import get_artifact_path, get_workspace_manifest, get_zone_path  # type: ignore[import-not-found]

# ── Chinese tokenizer (jieba) ──────────────────────────

_jieba = None
_USERDICT_LOADED = False


def _get_jieba_userdict_path() -> Path | None:
    """Return the path to the KOS jieba user dictionary, if bundled."""
    candidate = Path(__file__).with_suffix("").parent / "data" / "jieba_userdict.txt"
    return candidate if candidate.exists() else None


def _load_jieba_userdict() -> None:
    """Load the bundled user dictionary into jieba once per process."""
    global _USERDICT_LOADED
    if _USERDICT_LOADED:
        return
    dict_path = _get_jieba_userdict_path()
    if dict_path is None:
        _USERDICT_LOADED = True
        return
    try:
        import jieba

        jieba.load_userdict(str(dict_path))
    except Exception:
        pass
    _USERDICT_LOADED = True


def _tokenize_cn(text: str) -> str:
    """Tokenize Chinese text with jieba for FTS5 word-level indexing.
    Without this, FTS5 treats each Chinese char as a token — "数字化平台"
    can only be searched as individual characters, not as a phrase.
    With jieba, it becomes "数字化 平台" → both "数字化" and "平台" match.
    """
    global _jieba
    if _jieba is None:
        try:
            import jieba

            _jieba = jieba
            _load_jieba_userdict()
        except ImportError:
            return text  # jieba not installed, fallback to raw text
    # Tokenize: join words with spaces so FTS5 sees them as tokens
    tokens = _jieba.cut(text)
    return " ".join(tokens)


# ── Format Handlers ─────────────────────────────────────


class FormatHandler:
    extensions: set[str] = set()
    priority: int = 0

    def can_handle(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.extensions

    def extract_text(self, file_path: Path) -> str:
        raise NotImplementedError

    def extract_metadata(self, file_path: Path) -> dict:  # type: ignore[type-arg]
        return {"title": file_path.stem}


class MarkdownHandler(FormatHandler):
    extensions = {".md", ".txt", ".markdown", ".mdx"}
    priority = 100

    def extract_text(self, p: Path) -> str:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            return p.read_text(encoding="utf-8")[:8000]
        except Exception:  # noqa: BLE001
            return f"[{p.suffix}: {p.name}]"

    def extract_metadata(self, p: Path) -> dict:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            content = p.read_text(encoding="utf-8")
        except Exception:  # noqa: BLE001
            return {"title": p.stem}
        title = p.stem
        m = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
        if m:
            title = m.group(1).strip()
        # Extract KOS frontmatter comment
        fm = {}
        km = re.search(r"<!--\s*KOS:\s*(.+?)\s*-->", content[:500])
        if km:
            for part in km.group(1).split():
                if "=" in part:
                    k, v = part.split("=", 1)
                    fm[k.strip()] = v.strip()
        return {"title": title, **fm}


class DocxHandler(FormatHandler):
    extensions = {".docx"}
    priority = 80

    def extract_text(self, p: Path) -> str:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            from docx import Document

            doc = Document(str(p))
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())[:8000]
        except Exception:  # noqa: BLE001
            return f"[docx: {p.name}]"

    def extract_metadata(self, p: Path) -> dict:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            from docx import Document

            doc = Document(str(p))
            return {"title": doc.core_properties.title or p.stem}
        except Exception:  # noqa: BLE001
            return {"title": p.stem}


class GongwenDocxHandler(FormatHandler):
    """Structured docx extraction for gongwen domain: classify type, extract key sections."""

    extensions = {".docx"}
    priority = 85  # Higher than generic DocxHandler(80)

    def can_handle(self, p: Path) -> bool:  # type: ignore[reportIncompatibleMethodOverride]
        # Activate for gongwen and guozhuan zone docx files
        if p.suffix.lower() not in self.extensions:
            return False
        path_lower = str(p).lower()
        return any(kw in path_lower for kw in ["gongwen", "公文", "guozhuan", "国转中心"])

    # ── Type classification ──

    @staticmethod
    def _classify(paragraphs: list[dict]) -> str:  # type: ignore[type-arg]
        """Classify document type by content patterns."""
        full_text = " ".join(p["text"] for p in paragraphs[:30])
        # 通知
        if any(kw in full_text for kw in ["关于", "通知", "印发", "转发", "贯彻执行"]):
            if "特此通知" in full_text or "请遵照执行" in full_text or "请认真贯彻" in full_text:
                return "通知"
        # 方案
        if any(kw in full_text for kw in ["实施方案", "工作方案", "建设方案", "行动计划"]):
            return "方案"
        # 制度
        if any(kw in full_text for kw in ["管理办法", "管理规定", "实施细则", "制度", "章程", "条例"]):
            if "第一章" in full_text or "第一条" in full_text:
                return "制度"
        # 讲话稿
        if any(kw in full_text for kw in ["同志们", "各位领导", "讲话", "致辞", "汇报"]):
            return "讲话稿"
        # 报告
        if any(kw in full_text for kw in ["报告", "总结", "年度", "季度"]):
            return "报告"
        # 合同
        if any(kw in full_text for kw in ["合同", "协议", "甲方", "乙方", "签订"]):
            return "合同"
        return "通用"

    # ── Structured extraction ──

    @staticmethod
    def _extract_title(paragraphs: list[dict]) -> str:  # type: ignore[type-arg]
        """Find the title: first heading or first bold/large paragraph."""
        for p in paragraphs[:5]:
            if p.get("is_heading") or "Heading" in p.get("style", ""):
                return p["text"]  # type: ignore[no-any-return]
        for p in paragraphs[:3]:
            if len(p["text"]) > 5 and len(p["text"]) < 100:
                return p["text"]  # type: ignore[no-any-return]
        return ""

    @staticmethod
    def _extract_heading_structure(paragraphs: list[dict]) -> list[str]:  # type: ignore[type-arg]
        """Extract heading hierarchy for structure overview."""
        headings = []
        for p in paragraphs:
            style = p.get("style", "")
            if "Heading" in style or "标题" in style:
                level = 1
                for c in style:
                    if c.isdigit():
                        level = int(c)
                        break
                prefix = "#" * min(level, 4)
                headings.append(f"{prefix} {p['text']}")
        return headings[:20]  # Max 20 headings

    def _extract_notice(self, paragraphs: list[dict], doc) -> str:  # type: ignore[no-untyped-def, type-arg]
        """Extract structured content from 通知-type documents."""
        parts = []
        title = self._extract_title(paragraphs)
        if title:
            parts.append(f"# {title}")

        full_text = "\n".join(p["text"] for p in paragraphs)

        # Find 发文字号
        fwh = re.search(r"([一-鿿]+〔\d{4}〕\d+号)", full_text)
        if fwh:
            parts.append(f"发文字号: {fwh.group(1)}")

        # Main recipient 主送
        zs = re.search(r"(各[一-鿿]+、[一-鿿]+[：:])", full_text)
        if zs:
            parts.append(f"主送: {zs.group(1)}")

        # Body — first 2000 chars of non-heading text
        body_parts = []
        char_count = 0
        for p in paragraphs[3:]:  # Skip title + metadata lines
            if p.get("is_heading"):
                parts.append(f"\n## {p['text']}")
                continue
            text = p["text"]
            if char_count < 2000:
                body_parts.append(text)
                char_count += len(text)
            elif body_parts:
                break
        parts.extend(body_parts[:30])  # Max 30 paragraphs

        # Closing
        closing = re.search(r"(特此通知[。．.])", full_text)
        if closing:
            parts.append(closing.group(1))

        return "\n".join(parts)[:8000]

    def _extract_plan(self, paragraphs: list[dict], doc) -> str:  # type: ignore[no-untyped-def, type-arg]
        """Extract from 方案-type: 背景→目标→任务→分工."""
        parts = []
        title = self._extract_title(paragraphs)
        if title:
            parts.append(f"# {title}")

        # Extract sections by heading
        current_section = "正文"
        section_text = {"正文": []}  # type: ignore[var-annotated]

        for p in paragraphs:
            if p.get("is_heading"):
                current_section = p["text"][:30]
                section_text.setdefault(current_section, [])
            else:
                section_text.setdefault(current_section, []).append(p["text"])

        # Prioritize key sections
        key_sections = [
            s
            for s in section_text
            if any(
                kw in s
                for kw in [
                    "背景",
                    "目标",
                    "任务",
                    "分工",
                    "保障",
                    "要求",
                    "实施",
                    "组织",
                    "职责",
                    "指导思想",
                    "基本原则",
                    "工作",
                    "措施",
                    "步骤",
                    "进度",
                    "资金",
                ]
            )
        ]
        for s in key_sections or list(section_text.keys())[:8]:
            text = "\n".join(section_text[s][:5])
            if text.strip():
                parts.append(f"\n## {s}")
                parts.append(text[:500])

        return "\n".join(parts)[:8000]

    def _extract_regulation(self, paragraphs: list[dict], doc) -> str:  # type: ignore[no-untyped-def, type-arg]
        """Extract from 制度/规定: chapter→article structure."""
        parts = []
        title = self._extract_title(paragraphs)
        if title:
            parts.append(f"# {title}")

        chapters: list[dict[str, Any]] = []
        current_chapter: dict[str, Any] = {"title": "总则", "articles": []}

        for p in paragraphs:
            text = p["text"]
            ch_match = re.match(r"第[一二三四五六七八九十百千]+章\s*(.+)", text)
            if ch_match:
                if current_chapter["articles"]:
                    chapters.append(current_chapter)
                current_chapter = {"title": ch_match.group(1).strip(), "articles": []}
                continue

            art_match = re.match(r"(第[一二三四五六七八九十百千]+条)\s*(.+)", text)
            if art_match:
                current_chapter["articles"].append(f"{art_match.group(1)} {art_match.group(2)[:200]}")

        if current_chapter["articles"]:
            chapters.append(current_chapter)

        for ch in chapters[:8]:
            parts.append(f"\n## {ch['title']}")
            for art in ch["articles"][:5]:
                parts.append(art)

        return "\n".join(parts)[:8000]

    def _extract_speech(self, paragraphs: list[dict], doc) -> str:  # type: ignore[no-untyped-def, type-arg]
        """Extract from 讲话稿: title + first 2000 chars of body."""
        parts = []
        title = self._extract_title(paragraphs)
        if title:
            parts.append(f"# {title}")

        # Get first ~2000 chars of meaningful body text
        char_count = 0
        for p in paragraphs:
            if p.get("is_heading"):
                parts.append(f"\n## {p['text']}")
                continue
            text = p["text"]
            if len(text) > 10 and char_count < 2000:
                parts.append(text)
                char_count += len(text)
            elif char_count >= 2000:
                break

        return "\n".join(parts)[:8000]

    def _extract_generic(self, paragraphs: list[dict], doc) -> str:  # type: ignore[no-untyped-def, type-arg]
        """Generic extraction: title + headings + first paragraphs."""
        parts = []
        title = self._extract_title(paragraphs)
        if title:
            parts.append(f"# {title}")

        headings = self._extract_heading_structure(paragraphs)
        if headings:
            parts.append("\n## 文档结构")
            parts.extend(headings[:10])

        # First meaningful paragraphs
        char_count = 0
        for p in paragraphs:
            text = p["text"]
            if len(text) > 20 and char_count < 1500:
                parts.append(text)
                char_count += len(text)
            elif char_count >= 1500:
                break

        return "\n".join(parts)[:8000]

    def extract_text(self, p: Path) -> str:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            from docx import Document

            doc = Document(str(p))

            # Parse paragraph styles (max 500 paragraphs to prevent hang on huge docs)
            paragraphs: list[dict[str, Any]] = []
            for para in doc.paragraphs:
                if len(paragraphs) > 500:
                    break
                style = para.style.name if para.style else ""
                text = para.text.strip()
                if text:
                    paragraphs.append(
                        {
                            "style": style,
                            "text": text,
                            "is_heading": "Heading" in style or "标题" in style or "heading" in style.lower(),  # type: ignore[reportOptionalMemberAccess]
                        }
                    )

            if not paragraphs:
                return f"[docx: {p.name}]"

            # Classify document type
            doc_type = self._classify(paragraphs)

            # Extract by type
            if doc_type == "通知":
                text = self._extract_notice(paragraphs, doc)
            elif doc_type == "方案":
                text = self._extract_plan(paragraphs, doc)
            elif doc_type == "制度":
                text = self._extract_regulation(paragraphs, doc)
            elif doc_type == "讲话稿":
                text = self._extract_speech(paragraphs, doc)
            else:
                text = self._extract_generic(paragraphs, doc)

            return text
        except Exception:  # noqa: BLE001
            return f"[docx: {p.name}]"

    def extract_metadata(self, p: Path) -> dict:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            from docx import Document

            doc = Document(str(p))

            # Parse basic info
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(
                        {
                            "style": para.style.name if para.style else "",
                            "text": text,
                            "is_heading": "Heading" in (para.style.name or ""),  # type: ignore[union-attr]
                        }
                    )

            title = self._extract_title(paragraphs) if paragraphs else p.stem
            doc_type = self._classify(paragraphs) if paragraphs else "未知"

            meta = {"title": title or p.stem, "doc_type": doc_type}

            # Try extract 发文字号
            if paragraphs:
                full = "\n".join(p2["text"] for p2 in paragraphs[:10])  # type: ignore[misc]
                fwh = re.search(r"([一-鿿]+〔\d{4}〕\d+号)", full)
                if fwh:
                    meta["发文字号"] = fwh.group(1)

            return meta
        except Exception:  # noqa: BLE001
            return {"title": p.stem}


class PdfHandler(FormatHandler):
    extensions = {".pdf"}
    priority = 70

    def extract_text(self, p: Path) -> str:  # type: ignore[reportIncompatibleMethodOverride]
        # Tier 1: MinerU (best quality, Markdown output)
        try:
            from magic_pdf.tools.common import do_parse  # type: ignore[import-untyped, import-not-found]

            result = do_parse(str(p))
            md = result.get("markdown", "") or result.get("content", "")
            if md and len(md) > 50:
                return cast("str", md[:8000])
        except Exception:  # noqa: BLE001
            pass

        # Tier 2: PyMuPDF (fallback)
        try:
            import fitz  # type: ignore[import-untyped]

            doc = fitz.open(str(p))
            text = "\n".join(page.get_text() for page in doc)  # type: ignore[reportArgumentType]
            doc.close()
            return text[:8000]
        except Exception:  # noqa: BLE001
            return f"[pdf: {p.name}]"

    def extract_metadata(self, p: Path) -> dict:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            import fitz

            doc = fitz.open(str(p))
            meta = doc.metadata
            doc.close()
            return {"title": meta.get("title") or p.stem}  # type: ignore[reportOptionalMemberAccess]
        except Exception:  # noqa: BLE001
            return {"title": p.stem}


class XlsxHandler(FormatHandler):
    extensions = {".xlsx", ".xls"}
    priority = 60

    def extract_text(self, p: Path) -> str:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            import openpyxl  # type: ignore[unknown]

            wb = openpyxl.load_workbook(str(p), read_only=True)
            parts = []
            for sn in wb.sheetnames[:5]:
                ws = wb[sn]
                rows = []
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i > 50:
                        break
                    rows.append(" | ".join(str(c) if c is not None else "" for c in row))
                parts.append(f"Sheet: {sn}\n" + "\n".join(rows))
            wb.close()
            return "\n\n".join(parts)[:8000]
        except Exception:  # noqa: BLE001
            return f"[xlsx: {p.name}]"

    def extract_metadata(self, p: Path) -> dict:  # type: ignore[reportIncompatibleMethodOverride]
        try:
            import openpyxl

            wb = openpyxl.load_workbook(str(p), read_only=True)
            sheets = list(wb.sheetnames)
            wb.close()
            return {"title": p.stem, "sheets": sheets}
        except Exception:  # noqa: BLE001
            return {"title": p.stem}


class GenericHandler(FormatHandler):
    extensions = set()  # handles everything
    priority = 0

    def can_handle(self, p: Path) -> bool:  # type: ignore[reportIncompatibleMethodOverride]
        return True

    def extract_text(self, p: Path) -> str:  # type: ignore[reportIncompatibleMethodOverride]
        return f"[{p.suffix}: {p.name}]"

    def extract_metadata(self, p: Path) -> dict:  # type: ignore[reportIncompatibleMethodOverride]
        return {"title": p.name}


HANDLERS = [
    MarkdownHandler(),
    GongwenDocxHandler(),
    DocxHandler(),
    PdfHandler(),
    XlsxHandler(),
    GenericHandler(),
]


def check_handler_deps() -> dict:  # type: ignore[type-arg]
    """Check which handlers have their dependencies installed. Run before indexing."""
    status = {}
    # Markdown: always OK (stdlib)
    status["markdown"] = True
    # python-docx
    try:
        import docx

        status["docx"] = True
    except ImportError:
        status["docx"] = False
    # PyMuPDF
    try:
        import fitz

        status["pymupdf"] = True
    except ImportError:
        status["pymupdf"] = False
    # openpyxl
    try:
        import openpyxl

        status["openpyxl"] = True
    except ImportError:
        status["openpyxl"] = False
    # MinerU
    try:
        from magic_pdf.tools.common import do_parse  # type: ignore[reportMissingImports]

        status["mineru"] = True
    except ImportError:
        status["mineru"] = False
    # jieba
    try:
        import jieba

        status["jieba"] = True
    except ImportError:
        status["jieba"] = False
    return status


def _print_dep_status(status: dict):  # type: ignore[no-untyped-def, type-arg]
    """Print dependency status before indexing."""
    ok_list = [k for k, v in status.items() if v]
    missing = [k for k, v in status.items() if not v]
    print(f"  Dependencies: {' '.join(f'✅{d}' for d in ok_list)}", file=sys.stderr)
    if missing:
        print(f"  Missing: {' '.join(f'⚠️{d}' for d in missing)}", file=sys.stderr)
        for m in missing:
            pkg = {
                "docx": "python-docx",
                "pymupdf": "PyMuPDF",
                "openpyxl": "openpyxl",
                "mineru": "magic-pdf",
                "jieba": "jieba",
            }.get(m, m)
            print(f"    → pip install {pkg} --break-system-packages", file=sys.stderr)


def _verify_db_integrity(conn: Any) -> None:
    """Verify SQLite database integrity before indexing.
    Auto-repairs with VACUUM if corruption is detected.
    """
    cur = conn.execute("PRAGMA integrity_check")
    result = cur.fetchone()[0]
    if result != "ok":
        print(f"⚠️ Database integrity issue detected: {result}", file=sys.stderr, flush=True)
        print("   Attempting VACUUM repair...", file=sys.stderr, flush=True)
        try:
            conn.executescript("VACUUM; PRAGMA integrity_check;")
            cur = conn.execute("PRAGMA integrity_check")
            result = cur.fetchone()[0]
            if result == "ok":
                print("   ✅ VACUUM repair successful", file=sys.stderr, flush=True)
            else:
                print(f"   ❌ VACUUM failed: {result}", file=sys.stderr, flush=True)
        except Exception as exc:  # noqa: BLE001
            print(f"   ❌ Could not repair: {exc}", file=sys.stderr, flush=True)


def get_handler(file_path: Path) -> FormatHandler:
    candidates = [h for h in HANDLERS if h.can_handle(file_path)]
    return max(candidates, key=lambda h: h.priority)


# ── Indexer Core ─────────────────────────────────────────


class KosIndexer:
    def __init__(self) -> None:
        self.manifest = get_workspace_manifest()
        self.db_path = get_artifact_path("retrievalDatabase")
        self.now = datetime.now().strftime("%Y%m%d%H%M%S")

    def _connect(self) -> sqlite3.Connection:
        db_path = Path(str(self.db_path))
        db_path.parent.mkdir(parents=True, exist_ok=True)
        conn = sqlite3.connect(str(db_path))
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA busy_timeout=5000")
        conn.execute("PRAGMA synchronous=NORMAL")
        conn.execute("PRAGMA cache_size=-64000")  # 64MB cache
        conn.execute("PRAGMA temp_store=MEMORY")
        return conn

    def _batch_write(self, conn: sqlite3.Connection, pending: list[dict]) -> None:
        """Write a batch of processed files to the database in a single transaction."""
        cursor = conn.cursor()
        doc_cols = [
            "doc_id",
            "title",
            "kind",
            "zone",
            "status",
            "source",
            "owner",
            "created_at",
            "updated_at",
            "trust_level",
            "freshness",
            "review_status",
            "schema_version",
            "canonical_path",
            "source_url",
            "write_policy",
            "metadata_json",
            "body",
            "file_size",
            "file_mtime",
        ]
        doc_sql = f"INSERT INTO documents VALUES ({','.join(['?'] * len(doc_cols))})"
        fts_sql = "INSERT INTO documents_fts (doc_id,title,body,tags,canonical_path) VALUES (?,?,?,?,?)"
        fp_sql = """INSERT OR REPLACE INTO file_fingerprints
                    (canonical_path,zone,sha256_hash,file_size,file_mtime,last_indexed,absent_since,file_format)
                    VALUES (?,?,?,?,?,?,NULL,?)"""

        cursor.execute("BEGIN TRANSACTION")
        try:
            for r in pending:
                d = r["doc"]
                doc_id = d["doc_id"]
                # Delete old records first
                cursor.execute("DELETE FROM documents WHERE doc_id=?", (doc_id,))
                cursor.execute("DELETE FROM documents_fts WHERE doc_id=?", (doc_id,))
                # Insert new records
                cursor.execute(doc_sql, tuple(d[k] for k in doc_cols))
                cursor.execute(fts_sql, (d["doc_id"], d["title"], d["body"], "", d["canonical_path"]))
                fp = r["fingerprint"]
                cursor.execute(fp_sql, fp)
            conn.commit()
        except Exception:
            conn.rollback()
            raise

    def _init_schema(self, conn) -> None:  # type: ignore[no-untyped-def]
        conn.executescript("""
            CREATE TABLE IF NOT EXISTS documents (
                doc_id TEXT PRIMARY KEY,
                title TEXT,
                kind TEXT,
                zone TEXT,
                status TEXT,
                source TEXT,
                owner TEXT,
                created_at TEXT,
                updated_at TEXT,
                trust_level TEXT,
                freshness TEXT,
                review_status TEXT,
                schema_version TEXT,
                canonical_path TEXT,
                source_url TEXT,
                write_policy TEXT,
                metadata_json TEXT,
                body TEXT,
                file_size INTEGER,
                file_mtime TEXT
            );
            CREATE INDEX IF NOT EXISTS idx_documents_zone ON documents(zone);
            CREATE INDEX IF NOT EXISTS idx_documents_kind ON documents(kind);

            CREATE TABLE IF NOT EXISTS file_fingerprints (
                canonical_path TEXT PRIMARY KEY,
                zone TEXT NOT NULL,
                sha256_hash TEXT NOT NULL,
                file_size INTEGER NOT NULL,
                file_mtime TEXT NOT NULL,
                last_indexed TEXT NOT NULL,
                absent_since TEXT,
                file_format TEXT NOT NULL
            );
            CREATE INDEX IF NOT EXISTS idx_fp_zone ON file_fingerprints(zone);
            CREATE INDEX IF NOT EXISTS idx_fp_hash ON file_fingerprints(sha256_hash);
        """)
        # Ensure FTS table exists (standalone, not external content)
        conn.execute("""
            CREATE VIRTUAL TABLE IF NOT EXISTS documents_fts USING fts5(
                doc_id, title, body, tags, canonical_path
            );
        """)

    def _process_file(
        self,
        rel_path: str,
        abs_path: Path,
        zone_id: str,
        strategy: str,
        existing_hash: str | None = None,
    ) -> dict:  # type: ignore[type-arg]
        """Process a single file (stat + hash + extract). No DB access."""
        if strategy == "skip":
            return {"action": "skip", "rel": rel_path}
        try:
            fsize = abs_path.stat().st_size
            # Skip files > 100MB
            if fsize > 100 * 1024 * 1024:
                return {"action": "skip", "rel": rel_path}
            fmtime = datetime.fromtimestamp(abs_path.stat().st_mtime).strftime("%Y%m%d%H%M%S")

            if strategy == "filename_only":
                fhash = f"{fmtime}_{fsize}"
            else:
                fhash = self._compute_hash(abs_path)

            # Check if unchanged
            if existing_hash and existing_hash == fhash:
                return {"action": "skip", "rel": rel_path}  # shortcut

            # Build document record
            canonical = f"kos::{zone_id}::{rel_path}"
            doc_id = hashlib.sha1(canonical.encode()).hexdigest()
            handler = get_handler(abs_path)
            fmt = abs_path.suffix.lower().lstrip(".")

            if strategy == "filename_only":
                text = f"[{fmt}: {abs_path.name}]"
                meta = {"title": abs_path.name}
            else:
                text = handler.extract_text(abs_path)
                meta = handler.extract_metadata(abs_path)

            title = meta.get("title", abs_path.stem)
            body = text[:8000] if text else ""
            # Tokenize Chinese for FTS5 word-level matching
            body = _tokenize_cn(body)
            size = len(body.encode("utf-8"))

            return {
                "action": "added" if not existing_hash else "updated",
                "rel": rel_path,
                "doc": {
                    "doc_id": doc_id,
                    "title": title,
                    "kind": meta.get("kos_kind", "note"),
                    "zone": zone_id,
                    "status": meta.get("status", "active"),
                    "source": "auto-index",
                    "owner": "",
                    "created_at": self.now,
                    "updated_at": self.now,
                    "trust_level": "working",
                    "freshness": "active",
                    "review_status": "pending",
                    "schema_version": "1.0",
                    "canonical_path": canonical,
                    "source_url": "",
                    "write_policy": "managed",
                    "metadata_json": json.dumps(meta, ensure_ascii=False),
                    "body": body,
                    "file_size": size,
                    "file_mtime": fmtime,
                },
                "fingerprint": (rel_path, zone_id, fhash, fsize, fmtime, self.now, fmt),
            }
        except Exception as e:  # noqa: BLE001
            return {"action": "error", "rel": rel_path, "error": str(e)}

    def _compute_hash(self, file_path: Path) -> str:
        """SHA-256 of first 64KB + file size (fast, collision-resistant enough)."""
        h = hashlib.sha256()
        try:
            with open(file_path, "rb") as f:
                h.update(f.read(65536))
        except Exception:  # noqa: BLE001
            return "ERROR"
        h.update(str(file_path.stat().st_size).encode())
        return h.hexdigest()

    def _quick_unchanged_check(self, file_path: Path, existing_hash: str | None) -> bool:
        """Fast check if file is unchanged using mtime + size (avoids SHA-256).

        Returns True if the file is definitely unchanged, False if uncertain.
        """
        if not existing_hash:
            return False
        try:
            stat = file_path.stat()
            fmtime = datetime.fromtimestamp(stat.st_mtime).strftime("%Y%m%d%H%M%S")
            # Quick check: if mtime and size match, file is likely unchanged
            # The existing_hash contains SHA-256, but we can use mtime as a proxy
            # since the fingerprint also stores mtime
            return False  # Still need hash for certainty, but we can optimize later
        except Exception:
            return False

    def _get_indexing_strategy(self, zone_id: str, relative_path: str) -> str:
        """Determine indexing strategy for a file path within a zone."""
        strategies = self.manifest.get("indexing", {}).get("indexingStrategies", {}).get(zone_id, {})
        for prefix, strategy in strategies.items():
            if relative_path.startswith(prefix):
                return strategy  # type: ignore[no-any-return]
        return strategies.get("default", "full_text")  # type: ignore[no-any-return]

    def _get_scan_roots(self, zone_id: str, zone_config: dict) -> list[tuple[Path, str]]:
        """Resolve all scan roots for a zone (custom path + additionalPaths).

        Returns list of (resolved_path, label) tuples matching _iter_zone_files logic.
        """
        custom_path = zone_config.get("path")
        if custom_path:
            zone_path = Path(custom_path).expanduser().resolve()
        else:
            zone_path = Path(get_zone_path(zone_id))
        roots: list[tuple[Path, str]] = [(zone_path, str(zone_path))]
        for ap in zone_config.get("additionalPaths", []):
            ap_path = Path(ap).expanduser().resolve()
            if ap_path.exists():
                roots.append((ap_path, f"{zone_id}+{ap_path.name}"))
        return roots

    def _iter_zone_files(self, zone_id: str, zone_config: dict):  # type: ignore[no-untyped-def, type-arg]
        """Yield (relative_path, absolute_path) for all files in a zone.

        Supports:
          - followSymlinks: bool (default False) — follow symlinked directories
          - additionalPaths: [str] — secondary scan roots
        """
        scan_roots = self._get_scan_roots(zone_id, zone_config)
        global_excludes = self.manifest.get("indexing", {}).get("excludePrefixes", [])
        zone_excludes = zone_config.get("excludePrefixes", [])
        exclude_prefixes = list(set(global_excludes + zone_excludes))
        file_patterns = zone_config.get("filePatterns", [])
        follow_symlinks = zone_config.get("followSymlinks", False)

        def matches_file_patterns(rel_path: str) -> bool:
            if not file_patterns:
                return True
            filename = os.path.basename(rel_path)
            return any(
                fnmatch.fnmatch(rel_path, pattern) or fnmatch.fnmatch(filename, pattern) for pattern in file_patterns
            )

        for scan_root, root_label in scan_roots:
            if not Path(scan_root).exists():  # type: ignore[attr-defined]
                continue

            if follow_symlinks:
                # os.walk with followlinks=True for symlink support
                for dirpath, dirnames, filenames in os.walk(str(scan_root), followlinks=True):
                    # Skip hidden dirs
                    dirnames[:] = [d for d in dirnames if not d.startswith(".")]
                    for fname in filenames:
                        if fname.startswith("."):
                            continue
                        abs_path = Path(dirpath) / fname
                        if not abs_path.is_file():
                            continue
                        try:
                            rel = str(abs_path.relative_to(scan_root))
                        except ValueError:
                            continue
                        if not matches_file_patterns(rel):
                            continue
                        if any(rel.startswith(ep) for ep in exclude_prefixes):
                            continue
                        yield rel, abs_path
            else:
                # os.walk for normal (non-symlink) scanning — allows filtering
                # directories in-place to avoid OSError on deep trees (e.g. node_modules)
                for dirpath, dirnames, filenames in os.walk(str(scan_root)):
                    # Filter hidden dirs and excluded prefixes in-place
                    dirnames[:] = [
                        d
                        for d in dirnames
                        if not d.startswith(".")
                        and not any(d == ep or d.startswith(ep + os.sep) for ep in exclude_prefixes)
                    ]
                    rel_dir = os.path.relpath(dirpath, str(scan_root))
                    for fname in filenames:
                        if fname.startswith("."):
                            continue
                        abs_path = Path(dirpath) / fname
                        if not abs_path.is_file():
                            continue
                        rel = os.path.join(rel_dir, fname) if rel_dir != "." else fname
                        if not matches_file_patterns(rel):
                            continue
                        if any(rel.startswith(ep) or f"/{ep}/" in rel for ep in exclude_prefixes):
                            continue
                        yield rel, abs_path

    def index(self, incremental: bool = False, domain_filter: str | None = None, jobs: int = 4) -> dict[str, Any]:
        # Check dependencies first — prevent silent handler failures
        _print_dep_status(check_handler_deps())
        conn = self._connect()  # type: ignore[no-untyped-call]
        self._init_schema(conn)  # type: ignore[no-untyped-call]

        # Pre-verify database integrity to avoid mid-run corruption crashes
        _verify_db_integrity(conn)  # type: ignore[no-untyped-call]

        # Pre-load all fingerprints for incremental mode
        fingerprints = {}
        if incremental:
            rows = conn.execute("SELECT canonical_path, sha256_hash FROM file_fingerprints").fetchall()
            fingerprints = {r["canonical_path"]: r["sha256_hash"] for r in rows}

        stats = {"added": 0, "updated": 0, "skipped": 0, "removed": 0, "errors": 0}
        indexed_zones = []

        for zone_id, zone_config in self.manifest.get("zones", {}).items():
            if not zone_config.get("authoritative") or not zone_config.get("indexable"):
                continue
            if domain_filter and zone_id != domain_filter:
                continue
            indexed_zones.append(zone_id)

            # Scan files
            print(f"\n📂 {zone_id} — scanning files...", file=sys.stderr, flush=True)
            file_list = list(self._iter_zone_files(zone_id, zone_config))
            total = len(file_list)
            print(f"   {total} files → processing with {jobs} workers", file=sys.stderr, flush=True)

            # Process files in parallel
            zone_stats = {"added": 0, "updated": 0, "skipped": 0, "errors": 0, "removed": 0}
            pending = []
            batch_size = 2000

            with ThreadPoolExecutor(max_workers=jobs) as executor:
                futures = {}
                for rel_path, abs_path in file_list:
                    strategy = self._get_indexing_strategy(zone_id, rel_path)
                    existing = fingerprints.get(rel_path) if incremental else None
                    fut = executor.submit(self._process_file, rel_path, abs_path, zone_id, strategy, existing)  # type: ignore[arg-type]
                    futures[fut] = (rel_path, abs_path)

                completed = 0
                for fut in as_completed(futures):
                    completed += 1
                    result = fut.result()
                    action = result["action"]

                    if action == "skip":
                        zone_stats["skipped"] += 1
                    elif action == "error":
                        zone_stats["errors"] += 1
                        if zone_stats["errors"] <= 3:
                            print(f"\n  ⚠️  {result['rel']}: {result['error']}", file=sys.stderr)
                    else:
                        pending.append(result)
                        zone_stats[action] += 1

                    # Progress (every 100 files for responsiveness)
                    if completed % 100 == 0 or completed == total:
                        pct = completed * 100 // total
                        print(
                            f"\r   {completed}/{total} ({pct}%) +{zone_stats['added']} ~{zone_stats['updated']} ={zone_stats['skipped']}  ",
                            end="",
                            file=sys.stderr,
                            flush=True,
                        )

                    # Batch-write to DB (use single transaction)
                    if len(pending) >= batch_size or (completed == total and pending):
                        self._batch_write(conn, pending)
                        pending.clear()

            print(
                f"\r   ✅ {zone_id}: +{zone_stats['added']} new, ~{zone_stats['updated']} updated, ={zone_stats['skipped']} skipped  ",
                file=sys.stderr,
                flush=True,
            )
            for k in stats:
                stats[k] += zone_stats[k]

        # Cleanup: remove docs for absent files (absent > 3 days)
        if incremental:
            three_days_ago = (datetime.now() - timedelta(days=3)).strftime("%Y%m%d%H%M%S")
            absent = conn.execute(
                "SELECT canonical_path, zone FROM file_fingerprints WHERE absent_since IS NOT NULL AND absent_since < ?",
                (three_days_ago,),
            ).fetchall()
            try:
                for row in absent:
                    doc_id = hashlib.sha1(f"kos::{row['zone']}::{row['canonical_path']}".encode()).hexdigest()
                    conn.execute("DELETE FROM documents WHERE doc_id=?", (doc_id,))
                    conn.execute("DELETE FROM documents_fts WHERE doc_id=?", (doc_id,))
                    conn.execute("DELETE FROM file_fingerprints WHERE canonical_path=?", (row["canonical_path"],))
                    stats["removed"] += 1
                conn.commit()
            except sqlite3.DatabaseError as exc:
                print(f"  ⚠️ Cleanup deferred (non-fatal): {exc}", file=sys.stderr, flush=True)

        conn.commit()

        # Mark files that exist in fingerprints but not on disk as absent
        # NOTE: only run during incremental mode — full rebuild shouldn't mark
        # files absent in the same run, and the path check must use the same
        # scan-root resolution as _iter_zone_files (not bare get_zone_path)
        # to avoid false absent marks for custom paths/additionalPaths.
        if incremental:
            for zone_id in indexed_zones:
                zone_config = self.manifest.get("zones", {}).get(zone_id, {})
                scan_roots = self._get_scan_roots(zone_id, zone_config)
                fps = conn.execute(
                    "SELECT canonical_path FROM file_fingerprints WHERE zone=? AND absent_since IS NULL", (zone_id,)
                ).fetchall()
                for fp in fps:
                    exists = any((root / fp["canonical_path"]).exists() for root, _label in scan_roots)
                    if not exists:
                        conn.execute(
                            "UPDATE file_fingerprints SET absent_since=? WHERE canonical_path=?",
                            (self.now, fp["canonical_path"]),
                        )

        conn.commit()
        conn.close()

        return {
            "zones_indexed": indexed_zones,
            "mode": "incremental" if incremental else "full",
            "stats": stats,
            "timestamp": self.now,
        }

    def status(self) -> dict[str, Any]:
        conn = self._connect()  # type: ignore[no-untyped-call]
        self._init_schema(conn)  # type: ignore[no-untyped-call]
        total = conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
        zones = conn.execute("SELECT zone, COUNT(*) as cnt FROM documents GROUP BY zone ORDER BY cnt DESC").fetchall()
        fp_total = conn.execute("SELECT COUNT(*) FROM file_fingerprints").fetchone()[0]
        fp_absent = conn.execute("SELECT COUNT(*) FROM file_fingerprints WHERE absent_since IS NOT NULL").fetchone()[0]
        conn.close()
        return {
            "total_documents": total,
            "total_fingerprints": fp_total,
            "absent_files": fp_absent,
            "by_zone": {z["zone"]: z["cnt"] for z in zones},
        }

    def diff(self) -> dict[str, Any]:
        conn = self._connect()  # type: ignore[no-untyped-call]
        self._init_schema(conn)  # type: ignore[no-untyped-call]
        changed = []
        for zone_id in self.manifest.get("zones", {}):
            zone_config = self.manifest.get("zones", {}).get(zone_id, {})
            scan_roots = self._get_scan_roots(zone_id, zone_config)
            if not scan_roots:
                continue
            for rel, abs_p in self._iter_zone_files(zone_id, self.manifest["zones"][zone_id]):
                fhash = self._compute_hash(abs_p)
                existing = conn.execute(
                    "SELECT sha256_hash FROM file_fingerprints WHERE canonical_path=?", (rel,)
                ).fetchone()
                if not existing:
                    changed.append({"status": "new", "path": f"{zone_id}::{rel}"})
                elif existing["sha256_hash"] != fhash:
                    changed.append({"status": "changed", "path": f"{zone_id}::{rel}"})
                if len(changed) >= 50:
                    break
            if len(changed) >= 50:
                break
        conn.close()
        return {"changed_files": changed, "count": len(changed)}


def main() -> None:  # type: ignore[unknown]
    cmd = sys.argv[1] if len(sys.argv) > 1 else "status"
    incremental = "--incremental" in sys.argv
    domain = None
    jobs = 4
    for i, a in enumerate(sys.argv):
        if a == "--domain" and i + 1 < len(sys.argv):
            domain = sys.argv[i + 1]
        if a == "--jobs" and i + 1 < len(sys.argv):
            try:
                jobs = int(sys.argv[i + 1])
            except ValueError:
                pass

    indexer = KosIndexer()  # type: ignore[no-untyped-call]

    if cmd == "index":
        result = indexer.index(incremental=incremental, domain_filter=domain, jobs=jobs)  # type: ignore[arg-type]
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif cmd == "status":
        print(json.dumps(indexer.status(), ensure_ascii=False, indent=2))  # type: ignore[no-untyped-call]
    elif cmd == "diff":
        print(json.dumps(indexer.diff(), ensure_ascii=False, indent=2))  # type: ignore[no-untyped-call]
    else:
        print(json.dumps({"error": f"Unknown command: {cmd}"}))


if __name__ == "__main__":
    main()  # type: ignore[no-untyped-call]
